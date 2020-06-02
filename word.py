#!/usr/bin/env python
# -*- coding: utf-8 -*-


# 没有需要下载  pip install openpyxl
import docx

# 打开工作簿对象

PATH = '/Users/zhangxin/Desktop/model_test/'

document = docx.Document(PATH + 'article.docx')

for p in document.paragraphs:
    print('------------' + p.text)
    print(p.runs[0].font.name)
    # print(p.style)
#
# # 插入一级标题
# document.add_heading('古诗词', level=0)  #插入标题
# # 添加段落
# p = document.add_paragraph('''
#         人生就是一场抵达，我们总以为来日方长，可来日并不方长，我们总是在向往明天，而忽略了一个个今天，我们总是在仰望天空，却忘记要走好脚下的路。
# ''',)
# # 插入二级标题
# document.add_heading('春夜喜雨', level=1, )
#
# # 插入段落
# document.add_paragraph('好雨知时节，当春乃发生。', style='ListNumber')
# document.add_paragraph('随风潜入夜，润物细无声。', style='ListNumber')
# document.add_paragraph('野径云俱黑，江船火独明。', style='ListNumber')
# document.add_paragraph('晓看红湿处，花重锦官城。', style='ListNumber')
# # 保存文档
# document.save(PATH + 'article.docx')
